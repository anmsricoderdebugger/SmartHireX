import requests
import uuid
import xml.etree.ElementTree as ET
import traceback
from bs4 import BeautifulSoup   
from functools import wraps
from flask import Flask, request, jsonify, render_template, send_file, redirect, url_for, session
from authlib.integrations.flask_client import OAuth
from fpdf import FPDF
import os, json, re, time, tempfile, subprocess, traceback, io
import vertexai
from vertexai.generative_models import GenerativeModel
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import google.auth


load_dotenv()

app = Flask(__name__)
app.config['PREFERRED_URL_SCHEME'] = 'https'
app.secret_key = os.getenv('FLASK_SECRET_KEY')

# ── Google OAuth Setup ──
ALLOWED_DOMAIN = 'intellismith.com'  # Only this domain can access
GOOGLE_CLIENT_ID = os.getenv('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.getenv('GOOGLE_CLIENT_SECRET')

oauth = OAuth(app)
google = oauth.register(
   name='google',
   client_id=GOOGLE_CLIENT_ID,
   client_secret=GOOGLE_CLIENT_SECRET,
   server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
   client_kwargs={'scope': 'openid email profile'}
)
# --- CONFIG ---
# PROJECT_ID = "project-01-test-488412"
# REGION = "us-central1"

PROJECT_ID = os.getenv("PROJECT_ID")
REGION = os.getenv("REGION")


try:
   credentials, project = google.auth.default()
   vertexai.init(project=PROJECT_ID, location=REGION, credentials=credentials)
except Exception:
   vertexai.init(project=PROJECT_ID, location=REGION)

MODELS = [
   GenerativeModel("gemini-2.5-flash"),
   GenerativeModel("gemini-2.5-pro"),
   GenerativeModel("gemini-2.0-flash"),
]

MAX_CV_CHARS = 8000
MAX_JD_CHARS = 5000

CEIPAL_EMAIL = os.getenv("CEIPAL_EMAIL")
CEIPAL_PASSWORD = os.getenv("CEIPAL_PASSWORD")
CEIPAL_API_KEY = os.getenv("CEIPAL_API_KEY")

# ══════════ UTILITIES ══════════
def extract_text_from_file(file):
    """
    For manual uploaded Flask files: request.files['cv'] / request.files['jd']
    """
    try:
        if not file:
            return ""

        filename = (getattr(file, "filename", "") or "").lower()
        file.seek(0)

        if filename.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            file_bytes = file.read()
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                return _extract_docx(tmp_path)
            finally:
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass

        elif filename.endswith(".doc"):
            file_bytes = file.read()
            with tempfile.TemporaryDirectory() as tmpdir:
                doc_path = os.path.join(tmpdir, "input.doc")
                with open(doc_path, "wb") as f:
                    f.write(file_bytes)

                return (
                    _try_libreoffice(doc_path, tmpdir)
                    or _try_antiword(doc_path)
                    or _try_binary_doc(file_bytes)
                )

        else:
            return ""

    except Exception as e:
        print("❌ Extraction error:", e)
        return ""

def _extract_docx(source):
   doc = DocxDocument(source)
   para_text = "\n".join([p.text for p in doc.paragraphs]).strip()
   table_lines = []
   seen = set()
   for table in doc.tables:
       for row in table.rows:
           cells = [c.text.strip() for c in row.cells if c.text.strip()]
           if cells:
               line = " | ".join(cells)
               if line not in seen:
                   seen.add(line)
                   table_lines.append(line)
   table_text = "\n".join(table_lines)
   if len(para_text) > 100:
       extra = [l for l in table_lines if l[:50].lower() not in para_text.lower()]
       return para_text + ("\n\n" + "\n".join(extra) if extra else "")
   return table_text or para_text

def _try_libreoffice(doc_path, tmpdir):
   try:
       subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', tmpdir], capture_output=True, timeout=30)
       docx_path = os.path.join(tmpdir, "input.docx")
       if os.path.exists(docx_path):
           text = _extract_docx(docx_path)
           if text and len(text.strip()) > 50: return text
   except: pass
   return ""

def _try_antiword(doc_path):
   try:
       r = subprocess.run(['antiword', doc_path], capture_output=True, text=True, timeout=15)
       if r.returncode == 0 and len(r.stdout.strip()) > 50: return r.stdout.strip()
   except: pass
   return ""

def _try_binary_doc(file_bytes):
   runs, current = [], []
   for i in range(0, len(file_bytes) - 1, 2):
       lo, hi = file_bytes[i], file_bytes[i + 1]
       if hi == 0 and (32 <= lo <= 126 or lo in (9, 10, 13)):
           current.append(chr(lo))
       else:
           if len(current) > 20: runs.append(''.join(current).strip())
           current = []
   if len(current) > 20: runs.append(''.join(current).strip())
   runs.sort(key=len, reverse=True)
   clean = [r for r in runs[:15] if sum(1 for c in r if c.isalpha()) / max(len(r), 1) > 0.4 and len(r) > 30]
   result = '\n'.join(clean)
   return result if len(result) > 50 else ""


def extract_text_from_path(file_path):
    """
    For CEIPAL downloaded files saved in /tmp/...
    """
    try:
        path = os.fspath(file_path)
        ext = os.path.splitext(path.lower())[1]

        if ext == ".pdf":
            reader = PdfReader(path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext == ".docx":
            return _extract_docx(path)

        elif ext == ".doc":
            with open(path, "rb") as f:
                file_bytes = f.read()

            tmpdir = os.path.dirname(path) or tempfile.gettempdir()
            return (
                _try_libreoffice(path, tmpdir)
                or _try_antiword(path)
                or _try_binary_doc(file_bytes)
            )

        else:
            print("❌ Unsupported file type:", path)
            return ""

    except Exception as e:
        print("❌ Extraction error:", e)
        return ""


# In app.py — replace the clean_json function

def clean_json(raw):
   # Remove markdown code fences if present
   raw = re.sub(r'```json\s*', '', raw)
   raw = re.sub(r'```\s*', '', raw)
  
   # Extract JSON object or array
   match = re.search(r'[\{\[].*[\}\]]', raw, re.DOTALL)
   if match:
       raw = match.group(0)
  
   # Fix common Gemini JSON issues
   raw = raw.replace('\n', ' ').replace('\r', '')
   raw = re.sub(r',\s*}', '}', raw)   # Remove trailing commas before }
   raw = re.sub(r',\s*]', ']', raw)   # Remove trailing commas before ]
   raw = re.sub(r'[\x00-\x1f]', ' ', raw)  # Remove control characters
  
   try:
       data = json.loads(raw)
   except json.JSONDecodeError:
       # Last resort: try to find just the first complete JSON object
       depth = 0
       start = raw.index('{')
       for i in range(start, len(raw)):
           if raw[i] == '{': depth += 1
           elif raw[i] == '}': depth -= 1
           if depth == 0:
               data = json.loads(raw[start:i+1])
               break
       else:
           raise ValueError("Could not parse JSON from AI response")
  
   return data[0] if isinstance(data, list) else data

# In app.py — add this function after clean_json

def clean_json_flexible(raw):
   """More aggressive JSON extraction for Gemini responses that often break"""
   raw = re.sub(r'```json\s*', '', raw)
   raw = re.sub(r'```\s*', '', raw)
   raw = raw.strip()
  
   # Try direct parse first
   try:
       return json.loads(raw)
   except: pass
  
   # Clean and retry
   cleaned = raw.replace('\n', ' ').replace('\r', '')
   cleaned = re.sub(r',\s*}', '}', cleaned)
   cleaned = re.sub(r',\s*]', ']', cleaned)
   cleaned = re.sub(r'[\x00-\x1f]', ' ', cleaned)
  
   try:
       return json.loads(cleaned)
   except: pass
  
   # Try extracting array
   arr_match = re.search(r'\[.*\]', cleaned, re.DOTALL)
   if arr_match:
       try:
           return json.loads(arr_match.group(0))
       except: pass
  
   # Try extracting object
   obj_match = re.search(r'\{.*\}', cleaned, re.DOTALL)
   if obj_match:
       try:
           return json.loads(obj_match.group(0))
       except: pass
  
   # Last resort: extract individual question objects via regex
   questions = []
   pattern = r'\{"question"\s*:\s*"([^"]+)"\s*,\s*"options"\s*:\s*\[([^\]]+)\]\s*\}'
   for m in re.finditer(pattern, cleaned):
       opts = [o.strip().strip('"').strip("'") for o in m.group(2).split(',')]
       questions.append({"question": m.group(1), "options": opts[:4]})
  
   if questions:
       return questions
  
   raise ValueError("Could not parse AI response as JSON")

def call_gemini(prompt, retries=2):
   for model in MODELS:
       for attempt in range(retries):
           try:
               return model.generate_content(prompt)
           except Exception as e:
               if any(k in str(e).lower() for k in ["503", "429", "unavailable", "exhausted"]):
                   time.sleep((attempt + 1) * 3)
               else: raise
   raise Exception("All models failed")


def build_prompt(jd_text, cv_text, notes, candidate_id):

   return f"""
       Act as a Strategic Talent Architect. Conduct a forensic audit of the CV against the JD.

       Identify candidate name from CV. If missing use Candidate_{candidate_id}.

       Recruiter Notes:
       {notes or "None"}

       Rules:
       - Each point: Label: Short description
       - Max 3 items
       - No markdown

       Return ONLY JSON:

       {{
           "candidate_name": "",
           "overallScore": 0,
           "recommendation": "",
           "rationale": "",

           "strengths": {{
               "NIRF_and_Pedigree": [],
               "Experience_Alignment": [],
               "Projects_and_Quantifiable_Impact": []
           }},

           "proximity_matches": [],

           "gaps": {{
               "Functional_Gaps": [],
               "Domain_Mismatch": []
           }},

           "jd_enhancement": {{
               "missing_in_jd": []
           }}
       }}

       JD:
       {jd_text}

       CV:
       {cv_text}
   """



def clean_html_to_text(html):
   soup = BeautifulSoup(html, "html.parser")
   return soup.get_text(separator="\n")


# ── Auth Helper ──
def login_required(f):
   @wraps(f)
   def decorated(*args, **kwargs):
       user = session.get('user')
       if not user:
           return redirect('/login')
       return f(*args, **kwargs)
   return decorated

# ══════════ ROUTES ══════════

# ── Login Page ──
@app.route('/login')
def login_page():
   return '''
       <!DOCTYPE html>
       <html>
       <head>
           <title>Login - Smart AI Recruitment</title>
           <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
           <style>
               body { font-family: 'Inter', sans-serif; background: #f8fafc; display: flex;
                   align-items: center; justify-content: center; height: 100vh; margin: 0; }
               .login-box { background: white; padding: 40px; border-radius: 16px; text-align: center;
                           box-shadow: 0 4px 20px rgba(0,0,0,0.08); max-width: 400px; width: 90%; }
               .login-box h1 { font-size: 22px; margin-bottom: 8px; color: #0f172a; }
               .login-box p { color: #64748b; font-size: 14px; margin-bottom: 24px; }
               .google-btn { display: inline-flex; align-items: center; gap: 10px; padding: 12px 28px;
                           background: #2563eb; color: white; border: none; border-radius: 8px;
                           font-size: 15px; font-weight: 600; cursor: pointer; text-decoration: none; }
               .google-btn:hover { background: #1d4ed8; }
               .domain-note { margin-top: 16px; font-size: 12px; color: #94a3b8; }
           </style>
       </head>
       <body>
           <div class="login-box">
               <div style="font-size:40px; margin-bottom:12px;">🧠</div>
               <h1>Smart AI Recruitment</h1>
               <p>Sign in with your company Google account to continue</p>
               <a href="/auth/login" class="google-btn">🔐 Sign in with Google</a>
               <div class="domain-note">Only @intellismith.com accounts are allowed</div>
           </div>
       </body>
       </html>
   '''

# ── Google OAuth Flow ──
@app.route('/auth/login')
def auth_login():
    if app.debug:
        redirect_uri = url_for('auth_callback', _external=True)
        print("Redirect URI:", redirect_uri)
    else:
        redirect_uri = url_for('auth_callback', _external=True, _scheme='https')
        print("Redirect URI:", redirect_uri)
    return google.authorize_redirect(redirect_uri)

@app.route('/auth/callback')
def auth_callback():
   try:
       token = google.authorize_access_token()
       user_info = token.get('userinfo')
      
       if not user_info:
           return redirect('/login?error=no_info')
      
       email = user_info.get('email', '')
       domain = email.split('@')[-1] if '@' in email else ''
      
       # Domain restriction
       if domain != ALLOWED_DOMAIN:
           return f'''
           <html><body style="font-family:Inter,sans-serif; display:flex; align-items:center;
                 justify-content:center; height:100vh; background:#f8fafc;">
               <div style="text-align:center; background:white; padding:40px; border-radius:16px;
                    box-shadow:0 4px 20px rgba(0,0,0,0.08);">
                   <div style="font-size:48px; margin-bottom:12px;">🚫</div>
                   <h2 style="color:#dc2626;">Access Denied</h2>
                   <p style="color:#64748b;">Only @{ALLOWED_DOMAIN} accounts are allowed.</p>
                   <p style="color:#94a3b8; font-size:13px;">You signed in as: {email}</p>
                   <a href="/login" style="color:#2563eb; text-decoration:none; font-weight:600;">← Try again</a>
               </div>
           </body></html>
           '''
      
       # Store user in session
       session['user'] = {
           'email': email,
           'name': user_info.get('name', ''),
           'picture': user_info.get('picture', '')
       }
      
       return redirect('/')
   except Exception as e:
       print(f"Auth Error: {e}")
       return redirect('/login?error=auth_failed')

@app.route('/auth/logout')
def auth_logout():
   session.clear()
   return redirect('/login')


@app.route('/')
@login_required
def index():
   return render_template('index.html', user=session.get('user'))


# --- JD CREATOR: Generate MCQs from uploaded JD ---
@app.route('/api/jd/generate-mcqs', methods=['POST'])
@login_required
def generate_mcqs():
   try:
       jd_file = request.files.get('jd_file')
       if not jd_file: return jsonify({"success": False, "error": "No file uploaded"}), 400
       jd_text = extract_text_from_file(jd_file)[:MAX_JD_CHARS]
       if len(jd_text.strip()) == 0:
           return jsonify({"success": False, "error": "The uploaded JD file is blank. Please upload a valid document."}), 400
       if len(jd_text.strip()) < 50:
           return jsonify({"success": False, "error": "The uploaded JD appears to contain only images. Please upload a text-based document."}), 400

       prompt = f"""Analyze this job description and generate exactly 10 MCQ questions that will help refine and improve it.
       Each question should target a specific aspect: role clarity, required skills, experience level, compensation, work mode, team structure, growth path, culture fit, tools/tech, and diversity.
      
       Return ONLY valid JSON: {{"questions": [{{"question": "...", "options": ["A", "B", "C", "D"]}}]}}
      
       No markdown. No extra text. Exactly 10 questions, 4 options each.
      
       JD: {jd_text}"""

       response = call_gemini(prompt)
       raw = response.text
       print(f"[MCQ RAW] {raw[:500]}")
       parsed = clean_json_flexible(raw)
       if isinstance(parsed, dict) and 'questions' in parsed:
           questions = parsed['questions']
       elif isinstance(parsed, list):
           questions = parsed
       else:
           questions = []
       return jsonify({"success": True, "questions": questions[:10]})
   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500


# --- JD CREATOR: Create JD from MCQ answers ---
@app.route('/api/jd/create-from-mcqs', methods=['POST'])
@login_required
def create_from_mcqs():
   try:
       jd_file = request.files.get('jd_file')
       answers = json.loads(request.form.get('answers', '{}'))
       jd_text = extract_text_from_file(jd_file)[:MAX_JD_CHARS] if jd_file else ""

       if len(jd_text.strip()) == 0:
           return jsonify({"success": False, "error": "The uploaded JD file is blank. Please upload a valid document."}), 400
       if len(jd_text.strip()) < 50:
           return jsonify({"success": False, "error": "The uploaded JD appears to contain only images. Please upload a text-based document."}), 400

       prompt = f"""Create a professional, comprehensive job description based on the base JD and the recruiter's MCQ answers.

       FORMATTING: Return ONLY valid JSON: {{"jd_html": "<h3>SECTION</h3><p>Content</p>"}}
       Use <h3> for sections, <p> for paragraphs, <ul><li> for lists, <strong> for emphasis.
       No markdown. Make it professional and detailed.

       BASE JD: {jd_text}
       RECRUITER ANSWERS: {json.dumps(answers)}"""

       response = call_gemini(prompt)
       data = clean_json(response.text)
       return jsonify({"success": True, "jd_html": data.get("jd_html", "")})
   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500


# --- JD CREATOR: Create from manual input ---
@app.route('/api/jd/create-manual', methods=['POST'])
@login_required
def create_manual():
   try:
       payload = request.json
       prompt = f"""Create a professional, comprehensive job description from these details:

       Job Title: {payload.get('job_title', '')}
       Department: {payload.get('department', '')}
       Location: {payload.get('location', '')}
       Experience: {payload.get('experience', '')}
       Responsibilities: {payload.get('responsibilities', '')}
       Skills: {payload.get('skills', '')}
       Notes: {payload.get('notes', '')}

       FORMATTING: Return ONLY valid JSON: {{"jd_html": "<h3>SECTION</h3><p>Content</p>"}}
       Use <h3> for sections, <p> for paragraphs, <ul><li> for lists, <strong> for emphasis.
       Include: About the Role, Key Responsibilities, Requirements, Nice-to-Have, What We Offer.
       No markdown. Make it professional and detailed."""

       response = call_gemini(prompt)
       data = clean_json(response.text)
       return jsonify({"success": True, "jd_html": data.get("jd_html", "")})
   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500


# --- JD ENHANCER ---
@app.route('/api/jd/enhance', methods=['POST'])
@login_required
def enhance_jd():
   try:
       payload = request.json
       original = payload.get('original_jd', '')
       instructions = payload.get('instructions', '')

       prompt = f"""Enhance this job description based on the given instructions.
      
       Instructions: {instructions}
      
       FORMATTING: Return ONLY valid JSON: {{"enhanced_html": "<h3>SECTION</h3><p>Content</p>"}}
       Use <h3> for sections, <p> for paragraphs, <ul><li> for lists, <strong> for emphasis.
       Wrap enhanced/changed sections in <mark> tags so they are visually highlighted.
       No markdown.
      
       ORIGINAL JD: {original}"""

       response = call_gemini(prompt)
       data = clean_json(response.text)
       return jsonify({"success": True, "enhanced_html": data.get("enhanced_html", "")})
   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500

# In app.py — add this new route after the /api/jd/enhance route

# In app.py — replace the enhance_mcqs route

@app.route('/api/jd/enhance-mcqs', methods=['POST'])
@login_required
def enhance_mcqs():
   try:
       payload = request.json
       original_jd = payload.get('original_jd', '')[:3000]  # Limit input size

       prompt = f"""Generate 10 short MCQ questions to improve this job description.
       Topics: tone, skills detail, inclusivity, compensation, remote policy, growth path, team info, interview process, tech stack, DEI.
      
       STRICT RULES:
       - Each question max 15 words.
       - Each option max 8 words.
       - Exactly 4 options per question.
       - Return ONLY a JSON array, nothing else.
      
       FORMAT: [{{"question": "short question?", "options": ["A", "B", "C", "D"]}}]
      
       JD: {original_jd}"""

       response = call_gemini(prompt)
       raw = response.text
       print(f"[ENHANCE-MCQ RAW] {raw[:500]}")
      
       # Parse — could be array or object with questions key
       parsed = clean_json_flexible(raw)
      
       # Normalize: ensure it's a list of questions
       if isinstance(parsed, dict) and 'questions' in parsed:
           questions = parsed['questions']
       elif isinstance(parsed, list):
           questions = parsed
       else:
           questions = []
      
       return jsonify({"success": True, "questions": questions[:10]})
   except Exception as e:
       print(f"[ENHANCE-MCQ ERROR] {traceback.format_exc()}")
       return jsonify({"success": False, "error": str(e)}), 500

# --- JD DOWNLOAD AS DOCX ---
# In app.py — replace the entire download_docx function

@app.route('/api/jd/download-docx', methods=['POST'])
@login_required
def download_docx():
   try:
       html_content = request.json.get('html', '')
       doc = DocxDocument()
      
       # Parse HTML to DOCX
       text = html_content
       text = re.sub(r'<h3[^>]*>', '\n[HEADING]', text)
       text = re.sub(r'</h3>', '[/HEADING]\n', text)
       text = re.sub(r'<li[^>]*>', '\n• ', text)
       text = re.sub(r'<strong[^>]*>', '[B]', text)
       text = re.sub(r'</strong>', '[/B]', text)
       text = re.sub(r'<mark[^>]*>', '', text)
       text = re.sub(r'</mark>', '', text)
       text = re.sub(r'<[^>]+>', '', text)
       text = re.sub(r'\n{3,}', '\n\n', text).strip()

       for line in text.split('\n'):
           line = line.strip()
           if not line: continue
           if '[HEADING]' in line:
               heading_text = line.replace('[HEADING]', '').replace('[/HEADING]', '').replace('[B]', '').replace('[/B]', '').strip()
               doc.add_heading(heading_text, level=2)
           elif line.startswith('•'):
               item_text = line[1:].strip().replace('[B]', '').replace('[/B]', '')
               doc.add_paragraph(item_text, style='List Bullet')
           else:
               p = doc.add_paragraph()
               parts = re.split(r'(\[B\].*?\[/B\])', line)
               for part in parts:
                   if part.startswith('[B]'):
                       run = p.add_run(part.replace('[B]', '').replace('[/B]', ''))
                       run.bold = True
                   elif part.strip():
                       p.add_run(part)

       buffer = io.BytesIO()
       doc.save(buffer)
       buffer.seek(0)
       return send_file(buffer, as_attachment=True, download_name='Job_Description.docx',
                       mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   except Exception as e:
       print(f"DOCX Download Error: {traceback.format_exc()}")
       return jsonify({"success": False, "error": str(e)}), 500


# --- CV SCREENING MANUAL---
@app.route('/api/screen/manual', methods=['POST'])
@login_required
def screen_cv_manual():
   try:
       jd_file = request.files.get('jd')
       cv_file = request.files.get('cv')
       notes = request.form.get('notes', '')

       if not cv_file: return jsonify({"success": False, "error": "CV file missing"}), 400

       jd_text = extract_text_from_file(jd_file) if jd_file else "Not provided"
       cv_text = extract_text_from_file(cv_file)

       if not cv_text or len(cv_text.strip()) == 0:
           return jsonify({"success": False, "error": f"'{cv_file.filename}' appears to be a blank document."}), 400
       if len(cv_text.strip()) < 50:
           return jsonify({"success": False, "error": f"'{cv_file.filename}' appears to contain only images."}), 400

       jd_text = jd_text[:MAX_JD_CHARS]
       cv_text = cv_text[:MAX_CV_CHARS]

       prompt = f"""Act as a Strategic Talent Architect. Conduct a forensic audit of the CV against the JD.

           CORE INSTRUCTION: Identify the candidate's name from CV content or use the filename.

           RECRUITER OVERRIDES: {notes or "None."}

           FORMATTING RULES:
           - Every array item: "Label: Description in under 12 words"
           - Max 3 items per array. No "None" items — use empty array [].
           - No markdown. Rationale: one sentence, max 20 words.

           OUTPUT ONLY VALID JSON:
           {{
               "candidate_name": "Name", "overallScore": 0-100,
               "recommendation": "Level", "rationale": "One sentence.",
               "strengths": {{
                   "NIRF_and_Pedigree": [], "Experience_Alignment": [],
                   "Projects_and_Quantifiable_Impact": []
               }},
               "proximity_matches": [],
               "gaps": {{ "Functional_Gaps": [], "Domain_Mismatch": [] }},
               "jd_enhancement": {{ "missing_in_jd": [] }}
           }}

           JD: {jd_text}
           CV FILENAME: {cv_file.filename}
           CV: {cv_text}"""

       response = call_gemini(prompt)
       return jsonify({"success": True, "data": clean_json(response.text)})
   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500




##--------TOKEN AUTHORIZATION FOR CEIPAL ATS-----##

@login_required
def get_ceipal_token():
   print("\n[Ceipal Auth] Initiating token request...")
   url = 'https://api.ceipal.com/v1/createAuthtoken'
  
   payload = {
       "email": CEIPAL_EMAIL,
       "password": CEIPAL_PASSWORD,
       "api_key": CEIPAL_API_KEY
   }
   # We ask for JSON, but we will handle XML just in case
   headers = {'Content-Type': 'application/json', 'Accept': 'application/json'}
  
   try:
       response = requests.post(url, json=payload, headers=headers)
       print(f"[Ceipal Auth] HTTP Status Code: {response.status_code}")
       raw_body = response.text.strip()
       print(f"[Ceipal Auth] Raw Response Body: {raw_body}")

       if not raw_body:
           return None

       # --- XML HANDLING (Since your logs show XML) ---
       if raw_body.startswith('<?xml') or raw_body.startswith('<root'):
           print("[Ceipal Auth] Detected XML response. Parsing...")
           root = ET.fromstring(raw_body)
           # Find the access_token tag inside the XML
           token_element = root.find('access_token')
           if token_element is not None:
               token = token_element.text
               print(f"[Ceipal Auth] SUCCESS (XML)! Token: {token[:15]}...")
               return token

       # --- JSON HANDLING (Fallback) ---
       else:
           data = response.json()
           # Handle both list format and string format
           token_data = data.get('access_token')
           if isinstance(token_data, list) and len(token_data) > 0:
               token = token_data[0]
           else:
               token = token_data # assuming it's a string
          
           if token:
               print(f"[Ceipal Auth] SUCCESS (JSON)! Token: {token[:15]}...")
               return token

       print("[Ceipal Auth] FAILED: Could not find access_token in response.")
       return None
          
   except Exception as e:
       print(f"[Ceipal Auth] EXCEPTION: {str(e)}")
       return None

@login_required
def download_resume(url, access_token):
   try:
       headers = {
           "Authorization": f"Bearer {access_token}"
       }

       res = requests.get(url, headers=headers)

       if res.status_code != 200:
           print("❌ Download failed:", res.status_code)
           return None

       # ✅ Detect file type from headers
       content_type = res.headers.get("Content-Type", "").lower()
       print("Content-Type:", content_type)

       if "pdf" in content_type:
           ext = ".pdf"
       elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
           ext = ".docx"

       elif content_type == "application/msword":
           ext = ".doc"

       elif "word" in content_type:
           ext = ".docx"
      
       elif "word" in content_type or "docx" in content_type:
           ext = ".docx"
       else:
           # fallback (unknown type)
           ext = ".pdf"

       file_path = f"/tmp/{uuid.uuid4()}{ext}"

       with open(file_path, "wb") as f:
           f.write(res.content)

       print("✅ Saved:", file_path)

       return file_path

   except Exception as e:
       print("❌ Download error:", e)
       return None

@login_required
def get_ceipal_jobs(access_token, job_code):
   try:
       url = "https://api.ceipal.com/v1/getJobPostingsList"

       print(f"JOB CODE ------- {job_code}")
       params = {
           "searchkey": f'"{job_code}"'   # IMPORTANT format
       }

       headers = {
           "Content-Type": "application/json",
           "Authorization": f"Bearer {access_token}"
       }

       print("\n========== CEIPAL JOB API REQUEST ==========")
       print("URL:", url)
       print("Params:", params)
       print("Headers:", headers)

       response = requests.get(url, headers=headers, params=params)

       print("\n========== CEIPAL JOB API RESPONSE ==========")
       print("Status Code:", response.status_code)

       # ✅ RAW TEXT (exact response)
       print("\n--- RAW RESPONSE TEXT ---")
       print(response.text)

       # ✅ JSON Pretty Print (readable)
       try:
           data = response.json()
           print("\n--- FORMATTED JSON RESPONSE ---")
           print(json.dumps(data, indent=4))
           return data
       except Exception as json_err:
           print("\n❌ JSON Parsing Failed:", str(json_err))
           return None

   except Exception as e:
       print("\n❌ CEIPAL JOB FETCH ERROR:", str(e))
       return None


@login_required
def get_ceipal_submissions(access_token, job_id):
   """
   Fetches the list of submissions from Ceipal for a specific job.
   """
   print(f"\n[Ceipal Submissions] Requesting data for Job ID: {job_id}")
  
   # Based on your curl: the token is passed both in URL and Header
   url = f'https://api.ceipal.com/v1/getSubmissionsList?bearer%20token={access_token}&job_id={job_id}&isPipeline=1'
  
   headers = {
       'Content-Type': 'application/json',
       'Authorization': f'Bearer {access_token}'
   }
  
   try:
       response = requests.get(url, headers=headers)
       print(f"[Ceipal Submissions] HTTP Status Code: {response.status_code}")
      
       raw_body = response.text.strip()
       # Log a snippet of the response so it doesn't flood the terminal
       print(f"[Ceipal Submissions] Raw Response (First 200 chars): {raw_body[:200]}")

       if not raw_body:
           print("[Ceipal Submissions] ERROR: Received empty response.")
           return None
       else:
           data = response.json()
           print(f"[Ceipal Submissions] SUCCESS! Parsed JSON data.")
           print("---------------------")
           print(data)
           return data

       # Check if this API also returns XML like the Auth one did
       '''if raw_body.startswith('<?xml') or '<root>' in raw_body:
           print("[Ceipal Submissions] Parsing XML Data...")
           # If Ceipal returns candidates in XML, we'll need to parse tags here
           # For now, we return the raw text to see the structure
           return raw_body '''
      
      
   except Exception as e:
       print(f"[Ceipal Submissions] CRITICAL ERROR: {traceback.format_exc()}")
       return None


@app.route('/api/screen/ats', methods=['POST'])
@login_required
def screen_cv():
   try:
       data = request.get_json()

       notes = data.get('notes', '')
       job_code = data.get('job_code', '')

       # if not jd_text:
       #     return jsonify({"success": False, "error": "JD missing"}), 400

       if not job_code:
           return jsonify({"success": False, "error": "Job code missing"}), 400

       # ✅ 1. Get Token
       access_token = get_ceipal_token()
       if not access_token:
           return jsonify({"success": False, "error": "Token failed"}), 401

       # ✅ 2. Fetch Job using Job Code
       job_data = get_ceipal_jobs(access_token, job_code)
       job_info = job_data["results"][0]
       job_description_html = job_info.get("requisition_description", "")

       if not job_data or not job_data.get("results"):
           return jsonify({
               "success": False,
               "error": "No job found for given job code",
               "results": []
           }), 200

      
       # ✅ Extract job_id dynamically
       job_id = job_data["results"][0].get("id")
       if not job_id:
           return jsonify({"success": False, "error": "Job ID missing"}), 400

       print("✅ Job ID:", job_id)

       # ✅ 3. Fetch Submissions using dynamic job_id
       submissions = get_ceipal_submissions(access_token, job_id)
       print(submissions)

       if not submissions or "results" not in submissions:
           return jsonify({"success": False, "error": "No submissions found"}), 400

       candidates = submissions.get("results", [])
       print(f'Length of Candidates : {len(candidates)}')
       final_results = []

       for c in candidates:
           submission_id = c.get("submission_id")
           resume_url = c.get("resume")

           print(f"\nProcessing Candidate: {submission_id}")

           result = {
               "candidate_name": f"Candidate_{submission_id}",
               "overallScore": 0,
               "recommendation": "Failed",
               "rationale": "",
               "strengths": {},
               "proximity_matches": [],
               "gaps": {},
               "jd_enhancement": {}
           }

           if not resume_url:
               result["rationale"] = "No resume URL"
               final_results.append(result)
               continue

           file_path = download_resume(resume_url, access_token)

           if not file_path:
               result["rationale"] = "Download failed"
               final_results.append(result)
               continue

           cv_text = extract_text_from_path(file_path)
           jd_text = clean_html_to_text(job_description_html)

           if not cv_text or len(cv_text.strip()) < 50:
               result["rationale"] = "Unreadable CV"
               os.remove(file_path)
               final_results.append(result)
               continue

           try:
               prompt = build_prompt(jd_text, cv_text, notes, submission_id)
               response = call_gemini(prompt)
               parsed = clean_json(response.text)

               result.update(parsed)
               result["recommendation"] = result.get("recommendation", "Processed")

           except Exception as e:
               result["rationale"] = f"AI failed: {str(e)}"

           os.remove(file_path)
           final_results.append(result)
       print(f'final_results : {len(final_results)}')

       return jsonify({
           "success": True,
           "job_description_html": job_description_html,  # 👈 UI ke liye
           "results": final_results
       })

   except Exception as e:
       return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/cv/download-report', methods=['POST'])
@login_required
def download_cv_report():
   try:
       results = request.json.get('results', [])
       if not results:
           return jsonify({"success": False, "error": "No results"}), 400

       doc = DocxDocument()
      
       # Title
       title = doc.add_heading('Candidate Screening Report', level=1)
       doc.add_paragraph(f'Generated on {time.strftime("%d %B %Y")} • {len(results)} Candidates Screened')
       doc.add_paragraph('─' * 60)

       for d in results:
           score = d.get('overallScore', 0)
           name = d.get('candidate_name', 'Unknown')
           rec = d.get('recommendation', '')
           rationale = d.get('rationale', '')

           # Candidate header
           p = doc.add_heading(f'{name}  —  {score}%', level=2)
           doc.add_paragraph(f'Recommendation: {rec}')
           doc.add_paragraph(f'Rationale: {rationale}')

           # Strengths
           doc.add_heading('Strengths', level=3)
           for key in ['NIRF_and_Pedigree', 'Experience_Alignment', 'Projects_and_Quantifiable_Impact']:
               items = d.get('strengths', {}).get(key, [])
               for item in items:
                   if item and item.strip().lower() != 'none':
                       doc.add_paragraph(item, style='List Bullet')

           # Proximity
           doc.add_heading('Proximity', level=3)
           for item in d.get('proximity_matches', []):
               if item and item.strip().lower() != 'none':
                   doc.add_paragraph(item, style='List Bullet')

           # Gaps
           doc.add_heading('Gaps', level=3)
           for key in ['Functional_Gaps', 'Domain_Mismatch']:
               items = d.get('gaps', {}).get(key, [])
               for item in items:
                   if item and item.strip().lower() != 'none':
                       doc.add_paragraph(item, style='List Bullet')

           doc.add_paragraph('─' * 60)

       buffer = io.BytesIO()
       doc.save(buffer)
       buffer.seek(0)
       return send_file(buffer, as_attachment=True,
                       download_name=f'Candidate_Audit_Report_{time.strftime("%Y-%m-%d")}.docx',
                       mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   except Exception as e:
       print(f"Report Error: {traceback.format_exc()}")
       return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/jd/download-pdf', methods=['POST'])
@login_required
def download_pdf():
   try:
       html_content = request.json.get('html', '')
       print(f"[PDF DEBUG] HTML length: {len(html_content)}")
       print(f"[PDF DEBUG] First 200 chars: {html_content[:200]}")
       # Parse HTML to structured text
       text = html_content
       text = re.sub(r'<h3[^>]*>', '\n[HEADING]', text)
       text = re.sub(r'</h3>', '[/HEADING]\n', text)
       text = re.sub(r'<li[^>]*>', '\n[BULLET]', text)
       text = re.sub(r'<strong[^>]*>', '[B]', text)
       text = re.sub(r'</strong>', '[/B]', text)
       text = re.sub(r'<mark[^>]*>', '', text)
       text = re.sub(r'</mark>', '', text)
       text = re.sub(r'<[^>]+>', '', text)
       text = re.sub(r'\n{3,}', '\n\n', text).strip()
      
       # Sanitize unicode characters that Helvetica can't handle
       text = text.replace('\u2022', '-').replace('\u2013', '-').replace('\u2014', '-')
       text = text.replace('\u2018', "'").replace('\u2019', "'")
       text = text.replace('\u201c', '"').replace('\u201d', '"')
       text = text.replace('\u2026', '...').replace('\u00a0', ' ')
       text = text.replace('\u2023', '-').replace('\u25cf', '-').replace('\u25cb', '-')
       text = text.replace('\u2010', '-').replace('\u2011', '-').replace('\u2012', '-')
      
       pdf = FPDF()
       pdf.set_auto_page_break(auto=True, margin=20)
       pdf.add_page()
       pdf.set_font('Helvetica', '', 11)
      
       for line in text.split('\n'):
           line = line.strip()
           if not line:
               pdf.ln(4)
               continue
          
           if '[HEADING]' in line:
               heading = line.replace('[HEADING]', '').replace('[/HEADING]', '').replace('[B]', '').replace('[/B]', '').strip()
               pdf.ln(6)
               pdf.set_font('Helvetica', 'B', 14)
               pdf.cell(0, 8, heading, new_x="LMARGIN", new_y="NEXT")
               pdf.set_draw_color(37, 99, 235)
               pdf.line(10, pdf.get_y(), 200, pdf.get_y())
               pdf.ln(4)
               pdf.set_font('Helvetica', '', 11)
          
           elif '[BULLET]' in line:
               item = line.replace('[BULLET]', '').replace('[B]', '').replace('[/B]', '').strip()
               pdf.set_x(15)
               pdf.cell(5, 6, "-", new_x="END")
               pdf.multi_cell(170, 6, item, new_x="LMARGIN", new_y="NEXT")
          
           else:
               parts = re.split(r'(\[B\].*?\[/B\])', line)
               has_bold = any('[B]' in p for p in parts)
               if has_bold:
                   for part in parts:
                       if part.startswith('[B]'):
                           pdf.set_font('Helvetica', 'B', 11)
                           pdf.write(6, part.replace('[B]', '').replace('[/B]', ''))
                           pdf.set_font('Helvetica', '', 11)
                       elif part.strip():
                           pdf.write(6, part)
                   pdf.ln(6)
               else:
                   pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
      
       buffer = io.BytesIO()
       pdf.output(buffer)
       buffer.seek(0)
       return send_file(buffer, as_attachment=True, download_name='Job_Description.pdf',
                       mimetype='application/pdf')
   except Exception as e:
       print(f"PDF Error: {traceback.format_exc()}")
       return jsonify({"success": False, "error": str(e)}), 500

if __name__ == "__main__":
   app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
